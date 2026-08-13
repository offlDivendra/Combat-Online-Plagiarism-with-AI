"""
Text Extraction Utility
Extracts text from various file formats (TXT, PDF, DOCX).
"""

import os
from typing import Optional, Tuple
from pathlib import Path
import pypdf
from docx import Document


class TextExtractor:
    """
    Handles text extraction from multiple file formats.
    """
    
    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.docx'}
    
    def __init__(self):
        """Initialize the text extractor."""
        pass
    
    def is_supported(self, filename: str) -> bool:
        """
        Check if the file format is supported.
        
        Args:
            filename: Name of the file
            
        Returns:
            True if supported, False otherwise
        """
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS
    
    def get_file_extension(self, filename: str) -> str:
        """
        Get the file extension.
        
        Args:
            filename: Name of the file
            
        Returns:
            File extension (lowercase, with dot)
        """
        return Path(filename).suffix.lower()
    
    def extract_from_txt(self, file_path: str) -> str:
        """
        Extract text from a TXT file.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Extracted text
            
        Raises:
            FileNotFoundError: If file doesn't exist
            Exception: For other read errors
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text.strip()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                return text.strip()
            except Exception as e:
                raise Exception(f"Failed to read TXT file: {str(e)}")
        except FileNotFoundError:
            raise FileNotFoundError(f"File not found: {file_path}")
        except Exception as e:
            raise Exception(f"Error reading TXT file: {str(e)}")
    
    def extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text
            
        Raises:
            FileNotFoundError: If file doesn't exist
            Exception: For other read errors
        """
        try:
            text_parts = []
            
            with open(file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    raise Exception("PDF file is encrypted and cannot be read")
                
                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text:
                        text_parts.append(page_text)
            
            # Combine all pages
            full_text = '\n\n'.join(text_parts)
            return full_text.strip()
            
        except FileNotFoundError:
            raise FileNotFoundError(f"File not found: {file_path}")
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")
    
    def extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text
            
        Raises:
            FileNotFoundError: If file doesn't exist
            Exception: For other read errors
        """
        try:
            doc = Document(file_path)
            
            # Extract text from all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            # Combine all text
            full_text = '\n\n'.join(paragraphs)
            return full_text.strip()
            
        except FileNotFoundError:
            raise FileNotFoundError(f"File not found: {file_path}")
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
    
    def extract_text(self, file_path: str) -> Tuple[str, str]:
        """
        Extract text from a file based on its extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (extracted_text, file_extension)
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
            Exception: For other extraction errors
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = self.get_file_extension(file_path)
        
        if not self.is_supported(file_path):
            raise ValueError(
                f"Unsupported file format: {extension}. "
                f"Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        # Extract based on file type
        if extension == '.txt':
            text = self.extract_from_txt(file_path)
        elif extension == '.pdf':
            text = self.extract_from_pdf(file_path)
        elif extension == '.docx':
            text = self.extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
        
        # Validate extracted text
        if not text or not text.strip():
            raise Exception("No text could be extracted from the file")
        
        return text, extension
    
    def extract_with_metadata(self, file_path: str) -> dict:
        """
        Extract text along with file metadata.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with text and metadata
        """
        text, extension = self.extract_text(file_path)
        
        # Get file metadata
        file_stats = os.stat(file_path)
        filename = os.path.basename(file_path)
        
        return {
            "text": text,
            "filename": filename,
            "extension": extension,
            "file_size": file_stats.st_size,
            "word_count": len(text.split()),
            "character_count": len(text)
        }
    
    def batch_extract(self, file_paths: list) -> list:
        """
        Extract text from multiple files.
        
        Args:
            file_paths: List of file paths
            
        Returns:
            List of dictionaries with extraction results
        """
        results = []
        
        for file_path in file_paths:
            try:
                data = self.extract_with_metadata(file_path)
                data["status"] = "success"
                data["error"] = None
                results.append(data)
            except Exception as e:
                results.append({
                    "filename": os.path.basename(file_path),
                    "status": "error",
                    "error": str(e),
                    "text": None
                })
        
        return results
    
    def validate_file_size(self, file_path: str, max_size_mb: int = 10) -> bool:
        """
        Validate if file size is within acceptable limits.
        
        Args:
            file_path: Path to the file
            max_size_mb: Maximum file size in megabytes
            
        Returns:
            True if valid, False otherwise
        """
        try:
            file_size = os.path.getsize(file_path)
            max_size_bytes = max_size_mb * 1024 * 1024
            return file_size <= max_size_bytes
        except Exception:
            return False
    
    def get_file_info(self, file_path: str) -> dict:
        """
        Get basic file information without extracting text.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with file information
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_stats = os.stat(file_path)
        filename = os.path.basename(file_path)
        extension = self.get_file_extension(file_path)
        
        return {
            "filename": filename,
            "extension": extension,
            "file_size": file_stats.st_size,
            "file_size_mb": round(file_stats.st_size / (1024 * 1024), 2),
            "is_supported": self.is_supported(file_path),
            "created_at": file_stats.st_ctime,
            "modified_at": file_stats.st_mtime
        }


# Global extractor instance
text_extractor = TextExtractor()


# Convenience functions
def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from a file.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text
    """
    text, _ = text_extractor.extract_text(file_path)
    return text


def is_supported_file(filename: str) -> bool:
    """
    Check if a file format is supported.
    
    Args:
        filename: Name of the file
        
    Returns:
        True if supported, False otherwise
    """
    return text_extractor.is_supported(filename)


def get_supported_extensions() -> set:
    """
    Get set of supported file extensions.
    
    Returns:
        Set of supported extensions
    """
    return TextExtractor.SUPPORTED_EXTENSIONS
